from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from io import BytesIO
from PIL import Image, ImageEnhance
from docx.api import Document

class LeaveSpaceEnd:
    def __init__(self, value):
        self.value = value;

class LeaveSpaceInbetween:
    def __init__(self, value):
        self.value = value;

def copy_paragraphs(destination, source_paragraphs, leave_space, replace):
    for i, source_paragraph in enumerate(source_paragraphs):
        dest_paragraph = destination.add_paragraph();
        dest_paragraph.alignment = source_paragraph.alignment;
        dest_paragraph.paragraph_format.space_before = 0;

        if isinstance(leave_space, tuple):
            if i == len(source_paragraphs) - 1:
                dest_paragraph.paragraph_format.space_after = Pt(leave_space[1].value);
            else:
                dest_paragraph.paragraph_format.space_after = Pt(leave_space[0].value);

        elif isinstance(leave_space, LeaveSpaceInbetween):
            dest_paragraph.paragraph_format.space_after = Pt(leave_space.value);

        elif isinstance(leave_space, LeaveSpaceEnd) and i == len(source_paragraphs) - 1:
                dest_paragraph.paragraph_format.space_after = Pt(leave_space.value);

        else:
            dest_paragraph.paragraph_format.space_after = 0;

        # Iterate through the runs of the source paragraph and copy to the destination
        for run in source_paragraph.runs:
            new_run = None;
            for key in replace:
                if key in run.text.strip():
                    new_run = dest_paragraph.add_run(run.text.replace(key, replace[key]))
                    break;
            else:
                new_run = dest_paragraph.add_run(run.text)
            
            # Copy formatting
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.size = run.font.size
            new_run.font.name = run.font.name

def copy_header(DATA_FOLDER, dest_document, teacher_name, teacher_ci, teacher_type):
    header_document = Document(DATA_FOLDER + '/in/format/header.docx');
    fmt_section0 = header_document.sections[0];

    dest_section0 = dest_document.sections[0];
    dest_section0.different_first_page_header_footer = True;

    for rel in fmt_section0.first_page_header.part.rels.values():
        if "image" in rel.reltype and rel.target_ref:
            image_data = rel.target_part.blob;
            image = Image.open(BytesIO(image_data));

            image = image.convert("RGB");

            enhancer = ImageEnhance.Color(image)
            desaturated_image = enhancer.enhance(0.0)

            mod_image_stream = BytesIO()
            desaturated_image.save(mod_image_stream, format="PNG");

            dest_paragraph = dest_section0.first_page_header.add_paragraph()
            dest_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER;
            dest_paragraph.paragraph_format.space_before = 0;
            dest_paragraph.paragraph_format.space_after = 0;

            dest_run = dest_paragraph.add_run();
            dest_run.add_picture(mod_image_stream, width=Inches(0.8));

    copy_paragraphs(dest_section0.first_page_header, fmt_section0.first_page_header.paragraphs, LeaveSpaceEnd(10), {});

    copy_paragraphs(dest_document,
                    header_document.paragraphs,
                    (LeaveSpaceInbetween(10), LeaveSpaceEnd(20)),
                    {'$1PROFESOR': teacher_name, '$2CI': teacher_ci, '$3TIPO': teacher_type}
                    );

def copy_footer(DATA_FOLDER, dest_document, date):
    footer_document = Document(DATA_FOLDER + '/in/format/footer.docx')
    copy_paragraphs(dest_document,
                    footer_document.paragraphs,
                    None,
                    {'$1FECHA': date}
                    );

    footer_table = footer_document.tables[0];

    dest_table = dest_document.add_table(rows = 2, cols = 2);
    dest_table.style = 'Table Grid';

    for i, row in enumerate(footer_table.rows[0:]):
        dest_row = dest_table.rows[i].cells;
        for j, cell in enumerate(row.cells):
            dest_row[j].text = cell.text;

    for row in dest_table.rows:
        for cell in row.cells:
            cell.width = Pt(10);
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    font = run.font;
                    font.name = 'Arial';
                    font.size = Pt(9);