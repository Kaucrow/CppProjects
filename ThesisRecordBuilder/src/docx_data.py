import re
import os
from date import get_date
from tqdm import tqdm, trange
from docx.api import Document
from lxml import etree
import time

def get_docx_data_1(document_path, data, keys):
    if not hasattr(get_docx_data_1, 'thesis_count'):
        get_docx_data_1.thesis_count = 0

    document = Document(document_path)
    table = document.tables[0]

    filename = os.path.splitext(os.path.basename(document_path))[0]
    period = filename[filename.find('-') + 1: ]

    for i, row in enumerate(table.rows[1: ]):
        text = []

        for cell in row.cells:
            cleaned_cell_text = (re.sub(r' +', ' ',
                                    '\n'.join(
                                        line.strip() for line in cell.text.split('\n'))
                                    )
                                ).strip()
            text.append(cleaned_cell_text)
        
        text.append(period)

        # Construct a dictionary for this row, mapping
        # keys to values for this row
        row_data = dict(zip(keys, text))

        data.append(row_data)
        get_docx_data_1.thesis_count += 1

        #pbar.update(1)
    
    #pbar.close()

def get_docx_data_2(document_path, data, period, globals):
    if not hasattr(get_docx_data_2, 'thesis_count'):
        get_docx_data_2.thesis_count = 0

    document = Document(document_path)

    doc_xml = document.element.xml
    root = etree.fromstring(doc_xml)

    textboxes = root.xpath('//w:txbxContent', namespaces=root.nsmap)

    textbox_idx = 0

    paragraph_num = 0

    def check_idx(idx, delimiter):
        if idx == -1:
            raise Exception(f'Error in main paragraph of thesis in document (delimiter `{delimiter}` not found): {document_path}')

    thesis = {}

    for paragraph in document.paragraphs:
        paragraph = paragraph.text.strip()
        if not paragraph == '':
            paragraph_num += 1
            match paragraph_num:
                case 1:
                    if period in ['2017-C','2018-A','2018-B','2018-C','2019-A','2019-B','2019-C','2020-A','2020-B','2020-C','2021-A','2021-B','2021-C']:
                       paragraph_num = 0
                       continue

                    searchable = paragraph.lower()

                    # Thesis name
                    start_idx = searchable.find('especial:')
                    try:
                        check_idx(start_idx, 'especial:')
                        start_idx += len('especial:')
                    except Exception:
                        start_idx = searchable.find('grado:')
                        check_idx(start_idx, 'grado:')
                        start_idx += len('grado:')
                    end_idx = searchable.find('que', start_idx)
                    check_idx(end_idx, 'que')
                    thesis_name = paragraph[start_idx:end_idx].strip()
                    thesis['TITULO DE LA TESIS'] = thesis_name
                    searchable = searchable[end_idx:]
                    paragraph = paragraph[end_idx:]

                    # Student name
                    start_idx = searchable.find('bachiller:')
                    check_idx(start_idx, 'bachiller:')
                    start_idx += len('bachiller:')
                    end_idx = searchable.find('titular', start_idx)
                    check_idx(end_idx, 'titular')
                    student = paragraph[start_idx:end_idx].strip()
                    thesis['ALUMNO'] = student
                    searchable = searchable[end_idx:]
                    paragraph = paragraph[end_idx:]

                    # C.I.
                    found_with = ''
                    for delimiter in ['v-', 'v -', 'v –']:
                        start_idx = searchable.find(delimiter)
                        if start_idx != -1:
                            found_with = delimiter
                            break
                    check_idx(start_idx, 'v –')
                    start_idx += len(found_with)
                    substr = searchable[start_idx:]
                    match = re.search(r'[^0-9.]', substr.strip())
                    if match:
                        end_idx = match.start()
                    else:
                        check_idx(-1, None)
                    thesis['C.I.'] = substr[:end_idx + 1].strip()
                    searchable = searchable[end_idx:]
                    paragraph = paragraph[end_idx:]

                    # Grade
                    start_idx = searchable.find('(')
                    check_idx(start_idx, '(')
                    start_idx += len('(')
                    end_idx = searchable.find(')')
                    check_idx(end_idx, ')')
                    thesis['CALIFICACION'] = paragraph[start_idx:end_idx].strip()

                case 2:
                    # Thesis date
                    start_idx = paragraph.find('a los')
                    check_idx(start_idx, 'a los')
                    start_idx += len('a los')
                    end_idx = paragraph.find('.')
                    check_idx(end_idx, '.')
                    thesis['FECHA DE DEFENSA'] = get_date(paragraph[start_idx:end_idx].strip())

                    thesis['PERIODO'] = period

                    # Thesis teachers (Tutor and jury)
                    teachers = []

                    if period in ['2022-A','2022-B','2022-C']:
                        def check_title(title, expected):
                            if title != expected:
                                raise Exception(f'Error: teacher title ({title}) was not the expected title ({expected}) in thesis {document_path}')
                        def check_ci(ci):
                            if not ci[0].isdigit():
                               raise Exception(f'Error: found a bad teacher C.I. ({ci}) in thesis {document_path}')

                        table = document.tables[0]

                        try:
                            tutor_ci = table.rows[1].cells[2].text.strip()
                            check_ci(tutor_ci)
                        except Exception:
                            tutor_ci = table.rows[1].cells[3].text.strip()
                            check_ci(tutor_ci)
                        tutor_title = table.rows[2].cells[0].text.strip()
                        check_title(tutor_title, 'TUTOR')

                        try:
                            jurado_1_ci = table.rows[4].cells[1].text.strip()
                            check_ci(jurado_1_ci)
                        except Exception:
                            jurado_1_ci = table.rows[4].cells[2].text.strip()
                            check_ci(jurado_1_ci)
                        jurado_1_title = table.rows[5].cells[0].text.strip()
                        check_title(jurado_1_title, 'JURADO')

                        try:
                            jurado_2_ci = table.rows[4].cells[4].text.strip()
                            check_ci(jurado_2_ci)
                        except Exception:
                            try:
                                jurado_2_ci = table.rows[4].cells[5].text.strip()
                                check_ci(jurado_2_ci)
                            except Exception:
                                jurado_2_ci = table.rows[4].cells[6].text.strip()
                                check_ci(jurado_2_ci)
                        jurado_2_title = table.rows[5].cells[1].text.strip()
                        check_title(jurado_2_title, 'JURADO')

                        teachers = [tutor_ci, jurado_1_ci, jurado_2_ci]
                    else:
                        ci_texts = []
                        last_ci = ''
                        curr_ci = ''
                        for i in range(5):
                            textbox_texts = textboxes[textbox_idx + i].xpath('.//w:t/text()', namespaces=root.nsmap)
                            for text in textbox_texts:
                                if curr_ci:
                                    if text == 'TUTOR':
                                        teachers.insert(0, last_ci)
                                    elif text == 'JURADO':
                                        teachers.append(last_ci)

                                    last_ci = curr_ci
                                    curr_ci = ''

                                if globals.DEBUG:
                                    print(text)

                                if text[0].isdigit():
                                    ci_texts.append(text)
                                elif ci_texts and text[0] == '.':
                                    ci_texts.append(text)
                                elif ci_texts:
                                    curr_ci = ''.join(ci_texts)

                                    if last_ci == curr_ci:
                                        raise Exception('Found repeated textboxes')

                                    ci_texts = []

                        textbox_idx += 5

                    paragraph_num = 0

                    thesis['JURADO PRINCIPAL'] = teachers

                    data.append(thesis)

                    get_docx_data_2.thesis_count += 1

                    thesis = {}