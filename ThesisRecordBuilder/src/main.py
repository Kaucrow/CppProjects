import sys
import getopt
import os
import csv
from tqdm import tqdm, trange
from format_copy import *
from docx_data import *
from misc import num_to_str
from verdicts_data import get_verdicts_data, display
from docx.api import Document

def clear():
    if os.name == 'nt':
        _ = os.system('cls')
    else:
        _ = os.system('clear')

clear()

class Globals:
    VERSION: int = -1
    KNOWN_VERSIONS: int = [1, 2]
    DEBUG: bool = False
    DATA_FOLDER: str = 'data/'
    MISSING_STDTS: str = 'out/00_missingstudents.txt'
    KNOWN_MENTIONS = ['HONORIFICA', 'PUBLICACION']

global globals
globals = Globals()

runver_provided = False

try:
    opts, args = getopt.getopt(sys.argv[1:], 'ehdr:', ['example', 'help', 'debug', 'runver='])
except getopt.GetoptError as err:
    print(err)
    exit(1)

def print_versions():
        print('Available versions: ')
        for v in globals.KNOWN_VERSIONS:
            print(f'- v{v}')

for opt, arg in opts:
    if opt in ('-h', '--help'):
        print(
            '\n'.join(
                line.strip() for line in
                """
                [-h, --help]: Displays this help screen.
                [-d, --debug]: Runs the program in debug mode.
                [-e, --example]: Runs the program using the example files in the `data_ex` directory.
                """.split('\n')
            )
        )
        exit(0)

    if opt in ('-e', '--example'):
        globals.DATA_FOLDER = 'data_ex/'

    if opt in ('-d', '--debug'):
        globals.DEBUG = True

    if opt in ('-r', '--runver'):
        runver_provided = True
        try:
            version = int(arg)
            if version in globals.KNOWN_VERSIONS:
                globals.VERSION = version
            else:
                print(f'v{arg} is not a valid version.')
                print_versions()
                exit(1)
        except ValueError:
            print('The run version value must be an integer.')
            print_versions()
            exit(1)

if not runver_provided:
    print('Please pass the run version as an argument.')
    print_versions()
    exit(1)

print(f'Run version set to: {globals.VERSION}')

thesis_list = []
teachers_data = {}

if globals.VERSION == 1:
    try:
        os.remove(globals.DATA_FOLDER + globals.MISSING_STDTS)
    except OSError:
        pass

    teachers_names = {}

    with open(globals.DATA_FOLDER + 'in/teachers/teachers.csv', encoding='utf-8') as teachers_file:
        csv_reader = csv.reader(teachers_file, delimiter=';')
        line_count = 0
        for row in csv_reader:
            if line_count < 2:
                line_count += 1
                continue
            else:
                base_name = row[1].split('|')[0]
                for teacher in row[1].split('|'):
                    teachers_names[teacher] = base_name
                filename = base_name.replace(' ', '_').upper()
                teachers_data[base_name] = {'FILENAME':filename, 'FULL NAME':row[3] + ' ' + row[0], 'C.I.':row[2], 'THESIS':[]}

    keys = [
        'No', 'ALUMNO', 'C.I.', 'FECHA DE DEFENSA', 'TITULO DE LA TESIS', 'JURADO PRINCIPAL',
        'JURADO SUPLENTE', 'HORA', 'PERIODO', 'CALIFICACION', 'MENCION'
    ]

    calendars_path = globals.DATA_FOLDER + 'in/calendars/'
    calendar_files = os.listdir(calendars_path)

    print('Collecting data...')

    with tqdm(total=len(calendar_files), leave = True) as pbar_cal:
        for calendar in calendar_files:
            pbar_cal.set_description(f"Calendar: {calendar}")
            get_docx_data_1(calendars_path + calendar, thesis_list, keys)
            pbar_cal.update(1)

    verdicts_path = globals.DATA_FOLDER + 'in/verdicts/'
    verdict_files = os.listdir(verdicts_path)

    print()

    with tqdm(total=len(verdict_files), leave = True) as pbar_vrd:
        for verdict in verdict_files:
            pbar_vrd.set_description(f"Verdict: {os.path.basename(verdict)}")
            try:
                get_verdicts_data(verdicts_path + verdict, thesis_list, globals)
            except Exception as err:
                print('[ ERR ] Error occurred when processing verdict: ' + verdicts_path + verdict + '\n\terr: ' + err.args[0])
                print('\tTesseract text:\n\n' + err.args[3])
                display(err.args[1])
                display(err.args[2])
                exit(1)
            pbar_vrd.update(1)

    dest_document = Document();
    filename = '';
    curr_period = None

    for i in range(3):
        for idx, thesis in enumerate(thesis_list):
            teacher_found_name = thesis['JURADO PRINCIPAL'].split('\n')[i].replace('Arq.', '').strip()
            if teacher_found_name in teachers_names:
                teacher_real_name = teachers_names[teacher_found_name]
                if i == 0:
                    teachers_data[teacher_real_name]['THESIS'].append({'IDX':idx, 'TYPE':'tutor'})
                else:
                    teachers_data[teacher_real_name]['THESIS'].append({'IDX':idx, 'TYPE':'jury'})
            else:
                sys.exit('Could not find real name of teacher \'' + teacher_found_name + '\' (From thesis of period ' + thesis['PERIODO'] + ').\nPlease update the `teachers.csv` file')

    for teacher, info in teachers_data.items():
        info['THESIS'].sort(key = lambda x: thesis_list[x['IDX']]['PERIODO'])

elif globals.VERSION == 2:
    verdicts_path = globals.DATA_FOLDER + 'in/verdicts/'
    verdicts_folders = os.listdir(verdicts_path)
    for folder in verdicts_folders:
        verdicts_files = os.listdir(verdicts_path + folder)

        print(f'Collecting data (Period {folder})...')

        with tqdm(total=len(verdicts_files), leave = True) as pbar_vrd:
            period = folder
            for verdict in verdicts_files:
                pbar_vrd.set_description(f"Verdict: {verdict}")
                get_docx_data_2(verdicts_path + folder + '/' + verdict, thesis_list, period, globals)
                pbar_vrd.update(1)

        with open(globals.DATA_FOLDER + 'in/teachers/teachers.csv', encoding='utf-8') as teachers_file:
            csv_reader = csv.reader(teachers_file, delimiter=';')
            for i, row in enumerate(csv_reader):
                if i >= 2:
                    teacher_name = row[0]
                    ci = row[1]
                    title = row[2]
                    filename = teacher_name.replace(' ', '_').upper()
                    teachers_data[ci] = {'FILENAME':filename, 'FULL NAME':teacher_name, 'C.I.':ci, 'THESIS':[]}

        for idx, thesis in enumerate(thesis_list):
            for i, ci in enumerate(thesis['JURADO PRINCIPAL']):
                if ci not in teachers_data:
                    raise Exception(f'Error: could not find teacher with C.I. {ci} in teachers.csv, from thesis of student {thesis['ALUMNO']} from period {thesis['PERIODO']}')
                if i == 0:
                    teachers_data[ci]['THESIS'].append({'IDX':idx, 'TYPE':'tutor'})
                else:
                    teachers_data[ci]['THESIS'].append({'IDX':idx, 'TYPE':'jury'})

    for teacher, info in teachers_data.items():
        info['THESIS'].sort(key = lambda x: thesis_list[x['IDX']]['PERIODO'])

tutor_curr_period = ""
jury_curr_period = ""
tutor_doc = Document()
jury_doc = Document()
p = None

date = input("Elaboration date (for footer): ");

for teacher, teacher_info in teachers_data.items():
    tutor_curr_period = ""
    jury_curr_period = ""
    for thesis in teacher_info['THESIS']:
        print(thesis_list[thesis['IDX']]['PERIODO'])
        thesis_data = thesis_list[thesis['IDX']]
        if thesis['TYPE'] == 'tutor':
            if tutor_curr_period == "":
                tutor_doc = Document()
                copy_header(globals.DATA_FOLDER, tutor_doc, teacher_info['FULL NAME'], teacher_info['C.I.'], 'TUTOR')
            
            if thesis_data['PERIODO'] != tutor_curr_period:
                tutor_curr_period = thesis_data['PERIODO']
                p = tutor_doc.add_paragraph()
                p.style.font.name = 'Arial'
                p.style.font.size = Pt(11)
                p.paragraph_format.space_after = Pt(20)

                period_run = p.add_run('PERIODO ACADÉMICO ' + tutor_curr_period)
                period_run.bold = True
                period_run.italic = True

            p = tutor_doc.add_paragraph()

        elif thesis['TYPE'] == 'jury':
            if jury_curr_period == "":
                jury_doc = Document()
                copy_header(globals.DATA_FOLDER, jury_doc, teacher_info['FULL NAME'], teacher_info['C.I.'], 'JURADO')
            
            if thesis_data['PERIODO'] != jury_curr_period:
                jury_curr_period = thesis_data['PERIODO']
                p = jury_doc.add_paragraph()
                
                p.style.font.name = 'Arial'
                p.style.font.size = Pt(11)
                p.paragraph_format.space_after = Pt(20)

                period_run = p.add_run('PERIODO ACADÉMICO ' + jury_curr_period)
                period_run.bold = True
                period_run.italic = True
            
            p = jury_doc.add_paragraph()

        else:
            sys.exit('[ ERR ] Found unknown thesis type \'' + thesis['TYPE'] + '\' in teacher \'' + teacher + '\' data.')

        p.style.font.name = 'Arial'
        p.style.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(20)

        p.add_run('Nombre: ').italic = True
        name_run = p.add_run(thesis_data.get('ALUMNO').split('\n')[0].upper() + '\n')
        name_run.bold = True
        name_run.italic = True

        p.add_run('Titulo de T.E.G: ').italic = True
        title_run = p.add_run(' '.join(line.strip() for line in thesis_data.get('TITULO DE LA TESIS').split('\n')).upper() + '\n')
        title_run.bold = True
        title_run.italic = True

        if thesis_data.get('CALIFICACION'):
            p.add_run('Calificado: ').italic = True
            grade_run = p.add_run(num_to_str(thesis_data.get('CALIFICACION')).upper() + ' PUNTOS (' + thesis_data.get('CALIFICACION') + ')' + '\n')
            grade_run.bold = True
            grade_run.italic = True

        p.add_run('Fecha de Aprobación: ').italic = True
        date_run = p.add_run(thesis_data.get('FECHA DE DEFENSA'))
        date_run.bold = True
        date_run.italic = True

    if tutor_curr_period != "":
        copy_footer(globals.DATA_FOLDER, tutor_doc, date)
        tutor_doc.save(globals.DATA_FOLDER + 'out/CONSTANCIA_TUTOR_' + teacher_info['FILENAME'] + '.docx')

    if jury_curr_period != "":
        copy_footer(globals.DATA_FOLDER, jury_doc, date)
        jury_doc.save(globals.DATA_FOLDER + 'out/CONSTANCIA_JURADO_' + teacher_info['FILENAME'] + '.docx')

print("Finished execution.");